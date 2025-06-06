import os
from typing import List
from qdrant_client import QdrantClient
from qdrant_client.http.models import SearchParams
from openai import OpenAI
from dotenv import load_dotenv
import fitz
import docx
import io
load_dotenv()



from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract


# === OpenAI & Qdrant Configuration ===
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
QDRANT_URL = os.getenv("QDRANT_URL", "http://localhost:6333")
COLLECTION_EN = "uae_law_openai"
COLLECTION_AR = "uae_law_arabert"

qd_client = QdrantClient(
    url=f"https://{os.getenv('QDRANT_HOST')}",
    api_key=os.getenv("QDRANT_API_KEY"),
)


# === Custom Document Class ===
class Document:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata or {}


def extract_keywords(text: str) -> str:
    # Simple heuristic: return first 1–2 sentences as pseudo-query
    sentences = text.strip().split(".")
    return ". ".join(sentences[:2]).strip()

    
def extract_text(file_bytes: bytes) -> str:
    """Extracts and returns text from a PDF byte stream."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        doc.close()
        return full_text.strip()
    except Exception as e:
        return f"Error extracting text: {str(e)}"
    
# === Language Detection ===
def detect_language(text: str) -> str:
    arabic_chars = set("ابتثجحخدذرزسشصضطظعغفقكلمنهوي")
    if any(c in arabic_chars for c in text):
        return "ar"
    return "en"

# === OpenAI Embedding ===
def embed_text(text: str) -> List[float]:
    response = client.embeddings.create(
        model="text-embedding-3-small",  # or "text-embedding-ada-002"
        input=[text]
    )
    return response.data[0].embedding

# === Qdrant Semantic Search ===
def direct_qdrant_search(query: str, lang: str = "en", k: int = 10) -> List[Document]:
    embedding = embed_text(query)
    collection_name = COLLECTION_AR if lang == "ar" else COLLECTION_EN

    try:
        search_result = qd_client.search(
            collection_name=collection_name,
            query_vector=embedding,
            limit=k,
            search_params=SearchParams(hnsw_ef=128),
        )
    except Exception as e:
        print(f"[QDRANT ERROR] {e}")
        return []

    docs = []
    for hit in search_result:
        payload = hit.payload or {}
        content = payload.get("page_content") or payload.get("text", "")
        #print(f"[QDRANT HIT] {payload.get('source')} | Preview: {content[:100]}")
        metadata = {
            "source": payload.get("source", "unknown"),
            "page": payload.get("page", "N/A"),
        }
        docs.append(Document(page_content=content, metadata=metadata))

    return docs

def extract_text_from_upload(file_bytes):
    if not file_bytes:
        return ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return "\n".join([page.get_text() for page in doc])
    except Exception as e:
        print(f"[❌] PDF processing error: {e}")
        return ""

def extract_text_from_upload_all(filename, file_bytes):
    if filename.endswith(".pdf"):
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    elif filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join([para.text for para in doc.paragraphs])
    elif filename.endswith(".txt"):
        return file_bytes.decode("utf-8")
    else:
        raise ValueError("Unsupported file format. Please upload PDF, DOCX, or TXT.")


def extract_text_with_ocr(file_bytes, filename):
    # PDF OCR
    if filename.lower().endswith('.pdf'):
        try:
            images = convert_from_bytes(file_bytes)
            text = ''
            for img in images:
                text += pytesseract.image_to_string(img)
            return text.strip()
        except Exception as e:
            print(f"[OCR ERROR] Failed to OCR PDF {filename}: {e}")
            return ""

    # DOCX
    elif filename.lower().endswith('.docx'):
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            print(f"[DOCX ERROR] {filename}: {e}")
            return ""

    # Fallback: assume plain text
    else:
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            return ""